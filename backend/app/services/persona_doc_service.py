"""
═══════════════════════════════════════════════════════════════
PERSONA DOCUMENT SERVICE — Dynamic Loader for Persona .docx Files
═══════════════════════════════════════════════════════════════
Dynamically loads and pre-parses ALL domain-specific persona documents
from backend/app/data/personas_docs/ at application startup.

All content is served directly from the parsed documents — no GPT
prompts are used for question generation. The documents themselves
are the single source of truth.

Each persona doc is structured per-TASK with sections:
  - TASK + 5 Variants + 5 Adjacent Terms
  - SECTION 1 — Problems (5-8)
  - SECTION 2 — Opportunities (10-12)
  - SECTION 3 — Strategies (5-8)
  - SECTION 4 — RCA Bridge (5-8 lines)
"""

import os
import re
from functools import lru_cache
from pathlib import Path
from typing import Optional

import structlog

logger = structlog.get_logger()

# Path to the persona documents folder
PERSONAS_DOCS_DIR = Path(__file__).parent.parent / "data" / "personas_docs"

# ── Pre-loaded document cache ──────────────────────────────────
# Populated by preload_all_docs() at startup. Maps normalized domain → list of task blocks.
_DOC_CACHE: dict[str, list[dict]] = {}
_PRELOADED: bool = False


# ── Domain → Document Name Mapping ─────────────────────────────
# Maps the domain names (as used in the frontend) to the .docx file names.
# Keys are lowercase/normalized versions of domain names.

DOMAIN_TO_DOC: dict[str, str] = {
    "content & social media": "Content & Social Media.docx",
    "seo & organic visibility": "SEO & Organic Visibility.docx",
    "paid media & ads": "Paid Media & Ads.docx",
    "b2b lead generation": "B2B Lead Generation.docx",
    "sales execution & enablement": "Sales Execution & Enablement.docx",
    "lead management & conversion": "Lead Management & Conversion.docx",
    "customer success & reputation": "Customer Success & Reputation.docx",
    "repeat sales": "Same User More Sale_.docx",
    "repeate sales": "Same User More Sale_.docx",
    "business intelligence & analytics": "Business Intelligence & Analytics.docx",
    "market strategy & innovation": "Market Strategy & Innovation.docx",
    "financial health & risk": "Financial Health & Risk.docx",
    "org efficiency & hiring": "Org Efficiency & Hiring.docx",
    "improve yourself": "Owner_ Founder Improvements.docx",
    "sales & content automation": "Marketing  & Sales Automation.docx",
    "finance legal & admin": "Finance Legal & Admin.docx",
    "customer support ops": "Customer Support Ops.docx",
    "recruiting & hr ops": "Recruiting & HR Ops.docx",
    "personal & team productivity": "Personal & Team Productivity.docx",
    "marketing & sales automation": "Marketing  & Sales Automation.docx",
    "owner/founder improvements": "Owner_ Founder Improvements.docx",
    "same user more sale": "Same User More Sale_.docx",
}


def _extract_docx_text(filepath: Path) -> str:
    """Extract all text from a .docx file."""
    try:
        from docx import Document

        doc = Document(str(filepath))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("Failed to extract docx text", filepath=str(filepath), error=str(e))
        return ""


def _parse_task_blocks(full_text: str) -> list[dict]:
    """
    Parse the full document text into per-task blocks.

    Each task block starts with 'TASK:' and contains sections 1-4.
    Returns a list of dicts with keys:
      task, variants, adjacent_terms, problems, opportunities, strategies, rca_bridge, full_block
    """
    # Split on TASK: lines
    task_splits = re.split(r'(?=^TASK:\s)', full_text, flags=re.MULTILINE)
    blocks = []

    for block_text in task_splits:
        block_text = block_text.strip()
        if not block_text.startswith("TASK:"):
            continue

        parsed = {
            "task": "",
            "variants": "",
            "adjacent_terms": "",
            "problems": "",
            "opportunities": "",
            "strategies": "",
            "rca_bridge": "",
            "full_block": block_text,
        }

        # Extract task name
        task_match = re.match(r'TASK:\s*(.+)', block_text)
        if task_match:
            parsed["task"] = task_match.group(1).strip()

        # Extract variants
        variants_match = re.search(r'5 Variants:\n(.*?)(?=5 Adjacent Terms:)', block_text, re.DOTALL)
        if variants_match:
            parsed["variants"] = variants_match.group(1).strip()

        # Extract adjacent terms
        adj_match = re.search(r'5 Adjacent Terms:\n(.*?)(?=SECTION 1)', block_text, re.DOTALL)
        if adj_match:
            parsed["adjacent_terms"] = adj_match.group(1).strip()

        # Extract Section 1 — Problems
        s1_match = re.search(r'SECTION 1[^:]*:\n?(.*?)(?=SECTION 2)', block_text, re.DOTALL)
        if not s1_match:
            s1_match = re.search(r'SECTION 1.*?Problems.*?\n(.*?)(?=SECTION 2)', block_text, re.DOTALL)
        if s1_match:
            parsed["problems"] = s1_match.group(1).strip()

        # Extract Section 2 — Opportunities
        s2_match = re.search(r'SECTION 2[^:]*:\n?(.*?)(?=SECTION 3)', block_text, re.DOTALL)
        if not s2_match:
            s2_match = re.search(r'SECTION 2.*?Opportunities.*?\n(.*?)(?=SECTION 3)', block_text, re.DOTALL)
        if s2_match:
            parsed["opportunities"] = s2_match.group(1).strip()

        # Extract Section 3 — Strategies
        s3_match = re.search(r'SECTION 3[^:]*:\n?(.*?)(?=SECTION 4)', block_text, re.DOTALL)
        if not s3_match:
            s3_match = re.search(r'SECTION 3.*?Strategies.*?\n(.*?)(?=SECTION 4)', block_text, re.DOTALL)
        if s3_match:
            parsed["strategies"] = s3_match.group(1).strip()

        # Extract Section 4 — RCA Bridge
        s4_match = re.search(r'SECTION 4[^:]*:\n?(.*?)(?=TASK:|$)', block_text, re.DOTALL)
        if not s4_match:
            s4_match = re.search(r'SECTION 4.*?RCA.*?\n(.*?)(?=TASK:|$)', block_text, re.DOTALL)
        if s4_match:
            parsed["rca_bridge"] = s4_match.group(1).strip()

        blocks.append(parsed)

    return blocks


def _fuzzy_task_match(task_query: str, task_blocks: list[dict]) -> Optional[dict]:
    """
    Find the best matching task block for the user's selected task.
    Uses substring matching and normalized comparison.
    """
    query_lower = task_query.lower().strip()

    # Try exact match on task name first
    for block in task_blocks:
        if block["task"].lower().strip() == query_lower:
            return block

    # Try substring match (task query contained in task name or vice versa)
    for block in task_blocks:
        block_task = block["task"].lower().strip()
        if query_lower in block_task or block_task in query_lower:
            return block

    # Try matching against variants
    for block in task_blocks:
        variants_lower = block["variants"].lower()
        if query_lower in variants_lower:
            return block

    # Try word overlap scoring
    query_words = set(query_lower.split())
    best_score = 0
    best_block = None
    for block in task_blocks:
        block_words = set(block["task"].lower().split())
        overlap = len(query_words & block_words)
        if overlap > best_score:
            best_score = overlap
            best_block = block

    # Return best match if at least 2 words overlap
    if best_score >= 2 and best_block:
        return best_block

    # Fallback: return the first block
    if task_blocks:
        logger.warning(
            "No close task match found, using first block",
            query=task_query,
            available_tasks=[b["task"][:50] for b in task_blocks[:5]],
        )
        return task_blocks[0]

    return None


def _parse_rca_bridge_item(line: str) -> dict:
    """
    Parse a single RCA bridge line.
    Format: "symptom text" → metric → root_area
    Returns dict with: symptom, metric, root_area, raw
    """
    raw = line.strip()
    # Split on → (Unicode arrow)
    parts = raw.split("→")
    if len(parts) >= 3:
        return {
            "symptom": parts[0].strip().strip('\u201c\u201d"\''),
            "metric": parts[1].strip(),
            "root_area": parts[2].strip(),
            "raw": raw,
        }
    elif len(parts) == 2:
        return {
            "symptom": parts[0].strip().strip('\u201c\u201d"\''),
            "metric": parts[1].strip(),
            "root_area": "",
            "raw": raw,
        }
    return {"symptom": raw, "metric": "", "root_area": "", "raw": raw}


# ── Startup Pre-loader ─────────────────────────────────────────


def preload_all_docs():
    """
    Pre-parse ALL persona .docx files at application startup.
    Populates _DOC_CACHE so all subsequent lookups are instant.
    Called once from main.py lifespan.
    """
    global _PRELOADED
    if _PRELOADED:
        return

    # Parse each unique file once, then map domains to blocks
    file_to_blocks: dict[str, list[dict]] = {}
    for doc_name in set(DOMAIN_TO_DOC.values()):
        filepath = PERSONAS_DOCS_DIR / doc_name
        if filepath.exists():
            text = _extract_docx_text(filepath)
            blocks = _parse_task_blocks(text)
            file_to_blocks[doc_name] = blocks
            logger.info(
                "Parsed persona doc",
                file=doc_name,
                tasks=len(blocks),
                task_names=[b["task"][:50] for b in blocks],
            )
        else:
            logger.warning("Persona doc file not found", file=doc_name, path=str(filepath))

    # Map every domain key to its parsed blocks
    for domain_key, doc_name in DOMAIN_TO_DOC.items():
        if doc_name in file_to_blocks:
            _DOC_CACHE[domain_key] = file_to_blocks[doc_name]

    _PRELOADED = True
    logger.info(
        "All persona docs preloaded",
        domains_mapped=len(_DOC_CACHE),
        unique_files=len(file_to_blocks),
        total_tasks=sum(len(blocks) for blocks in file_to_blocks.values()),
    )


def _get_blocks_for_domain(domain: str) -> list[dict]:
    """
    Get parsed task blocks for a domain. Uses preloaded cache first,
    falls back to on-demand loading.
    """
    domain_lower = domain.lower().strip()

    # Exact cache hit
    if domain_lower in _DOC_CACHE:
        return _DOC_CACHE[domain_lower]

    # Fuzzy cache hit
    for key in _DOC_CACHE:
        if key in domain_lower or domain_lower in key:
            return _DOC_CACHE[key]

    # Fallback: load on demand (shouldn't happen after preload)
    doc_name = DOMAIN_TO_DOC.get(domain_lower)
    if not doc_name:
        for key, value in DOMAIN_TO_DOC.items():
            if key in domain_lower or domain_lower in key:
                doc_name = value
                break
    if doc_name:
        filepath = PERSONAS_DOCS_DIR / doc_name
        if filepath.exists():
            text = _extract_docx_text(filepath)
            blocks = _parse_task_blocks(text)
            _DOC_CACHE[domain_lower] = blocks
            logger.info("On-demand doc load", domain=domain, tasks=len(blocks))
            return blocks

    return []


# ── Public API: Dynamic Section Loader ─────────────────────────


def get_diagnostic_sections(domain: str, task: str) -> Optional[dict]:
    """
    Get structured diagnostic sections from the persona doc for a task.

    Returns the document's own content as structured sections — no GPT involved.
    This is the core dynamic loader: Problems, RCA Bridge, Opportunities
    are served directly from the .docx files.

    Returns:
        {
            "task_matched": "exact task name from doc",
            "sections": [
                {
                    "key": "problems",
                    "label": "Problem Areas",
                    "question": "Which of these problem areas best describes your current challenge?",
                    "items": ["Problem 1 text", "Problem 2 text", ...]
                },
                {
                    "key": "rca_bridge",
                    "label": "Diagnostic Signals",
                    "question": "Which of these symptoms are you experiencing?",
                    "items": ["Symptom 1", "Symptom 2", ...],
                    "rca_parsed": [{"symptom": ..., "metric": ..., "root_area": ...}, ...]
                },
                {
                    "key": "opportunities",
                    "label": "Growth Opportunities",
                    "question": "Which of these opportunities would be most valuable for your situation?",
                    "items": ["Opp 1 text", "Opp 2 text", ...]
                }
            ],
            "strategies": "raw strategies text",
            "full_context": { ... full parsed task block ... }
        }
    """
    blocks = _get_blocks_for_domain(domain)
    if not blocks:
        logger.warning("No task blocks for domain", domain=domain)
        return None

    matched = _fuzzy_task_match(task, blocks)
    if not matched:
        logger.warning("No matching task", domain=domain, task=task)
        return None

    sections = []

    # ── Section 1: Problems ────────────────────────────────────
    if matched["problems"]:
        items = [
            line.strip()
            for line in matched["problems"].split("\n")
            if line.strip() and len(line.strip()) > 15
        ]
        if items:
            sections.append({
                "key": "problems",
                "label": "Problem Areas",
                "question": "Which of these problem areas best describes your current challenge?",
                "items": items[:8],
                "allows_free_text": True,
            })

    # ── Section 4: RCA Bridge ──────────────────────────────────
    if matched["rca_bridge"]:
        raw_lines = [
            line.strip()
            for line in matched["rca_bridge"].split("\n")
            if line.strip() and len(line.strip()) > 15
        ]
        parsed_items = [_parse_rca_bridge_item(line) for line in raw_lines]
        # Show just the symptom part for user-facing options
        symptom_items = [item["symptom"] for item in parsed_items if item["symptom"]]
        if symptom_items:
            sections.append({
                "key": "rca_bridge",
                "label": "Diagnostic Signals",
                "question": "Which of these symptoms are you experiencing?",
                "items": symptom_items[:8],
                "allows_free_text": True,
                "rca_parsed": parsed_items,
            })

    # ── Section 2: Opportunities ───────────────────────────────
    if matched["opportunities"]:
        items = [
            line.strip()
            for line in matched["opportunities"].split("\n")
            if line.strip() and len(line.strip()) > 15
        ]
        if items:
            sections.append({
                "key": "opportunities",
                "label": "Growth Opportunities",
                "question": "Which of these opportunities would be most valuable for your situation?",
                "items": items[:8],
                "allows_free_text": True,
            })

    result = {
        "task_matched": matched["task"],
        "sections": sections,
        "strategies": matched.get("strategies", ""),
        "full_context": matched,
    }

    logger.info(
        "Diagnostic sections loaded from doc",
        domain=domain,
        task=task,
        matched_task=matched["task"][:60],
        num_sections=len(sections),
        section_sizes=[len(s["items"]) for s in sections],
    )

    return result


# ── Backward-compatible API ────────────────────────────────────


@lru_cache(maxsize=32)
def _load_raw_doc(domain_lower: str) -> str:
    """Load & cache the raw text of a persona doc by domain."""
    doc_name = DOMAIN_TO_DOC.get(domain_lower)
    if not doc_name:
        for key, value in DOMAIN_TO_DOC.items():
            if key in domain_lower or domain_lower in key:
                doc_name = value
                break
    if not doc_name:
        return ""
    filepath = PERSONAS_DOCS_DIR / doc_name
    if not filepath.exists():
        return ""
    return _extract_docx_text(filepath)


@lru_cache(maxsize=32)
def load_persona_doc(domain: str) -> Optional[str]:
    """Load the full raw persona document content for a given domain."""
    content = _load_raw_doc(domain.lower().strip())
    if content:
        return content
    return None


def load_task_context(domain: str, task: str) -> Optional[dict]:
    """
    Load the specific task context from a persona doc.
    Used by the recommendation engine (agent_service.py).
    """
    blocks = _get_blocks_for_domain(domain)
    if not blocks:
        return None
    matched = _fuzzy_task_match(task, blocks)
    return matched


def get_all_tasks_for_domain(domain: str) -> list[str]:
    """Return all task names found in a persona doc for a domain."""
    blocks = _get_blocks_for_domain(domain)
    return [b["task"] for b in blocks]


def get_available_personas() -> list[str]:
    """Return list of all domains that have persona documents available."""
    available = []
    for domain, doc_name in DOMAIN_TO_DOC.items():
        filepath = PERSONAS_DOCS_DIR / doc_name
        if filepath.exists():
            available.append(domain)
    return available


def get_doc_for_domain(domain: str) -> Optional[str]:
    """Get the persona document name for a domain."""
    domain_lower = domain.lower().strip()
    return DOMAIN_TO_DOC.get(domain_lower)
